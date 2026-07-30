"""
Document parser implementations.
Extracts text content from PDF, DOCX, XLSX, TXT, and HTML files.
"""

import structlog
from pathlib import Path

from src.modules.knowledge.domain.interfaces.repositories import IDocumentParser

logger = structlog.get_logger()


class DocumentParser(IDocumentParser):
    """
    Unified document parser that delegates to format-specific extractors.
    Supports: PDF, DOCX, XLSX, TXT, HTML.
    """

    MIME_MAP = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "text/plain": "txt",
        "text/html": "html",
        "text/csv": "txt",
        "text/markdown": "txt",
    }

    async def parse(self, file_path: str, mime_type: str) -> str:
        """Extract text from a file based on its MIME type."""
        format_key = self.MIME_MAP.get(mime_type)
        if not format_key:
            # Try to detect from extension
            ext = Path(file_path).suffix.lower().lstrip(".")
            format_key = ext if ext in ("pdf", "docx", "xlsx", "txt", "html") else None

        if not format_key:
            raise ValueError(f"Unsupported file type: {mime_type}")

        logger.info("Parsing document", file_path=file_path, format=format_key)

        if format_key == "pdf":
            return await self._parse_pdf(file_path)
        elif format_key == "docx":
            return await self._parse_docx(file_path)
        elif format_key == "xlsx":
            return await self._parse_xlsx(file_path)
        elif format_key == "html":
            return await self._parse_html(file_path)
        else:
            return await self._parse_txt(file_path)

    def supported_types(self) -> list[str]:
        return list(self.MIME_MAP.keys())

    async def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())

        return "\n\n".join(pages)

    async def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    async def _parse_xlsx(self, file_path: str) -> str:
        """Extract text from Excel using openpyxl."""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        content_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            content_parts.append(f"## Sheet: {sheet_name}")

            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) for cell in row if cell is not None]
                if cells:
                    content_parts.append(" | ".join(cells))

        wb.close()
        return "\n".join(content_parts)

    async def _parse_html(self, file_path: str) -> str:
        """Extract text from HTML using BeautifulSoup."""
        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        text = soup.get_text(separator="\n", strip=True)
        # Clean up multiple newlines
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n\n".join(lines)

    async def _parse_txt(self, file_path: str) -> str:
        """Read plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()


class WebPageParser:
    """Fetches and parses web pages for indexing."""

    async def fetch_and_parse(self, url: str) -> str:
        """Download a web page and extract its text content."""
        import httpx
        from bs4 import BeautifulSoup

        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.get(url)
            response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")

        # Remove non-content elements
        for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
            element.decompose()

        # Try to get main content
        main = soup.find("main") or soup.find("article") or soup.find("body")
        if not main:
            main = soup

        text = main.get_text(separator="\n", strip=True)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n\n".join(lines)
